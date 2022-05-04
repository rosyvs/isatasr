import re
import string
import docx
import Levenshtein
import pandas as pd
import numpy as np
import collections
import contextlib
from pathlib import Path
import os
import wave
from num2words import num2words
from decimal import InvalidOperation
from pydub import AudioSegment
from google.cloud import speech
from google.cloud import speech_v1p1beta1 as speechB # need the beta for diarizaiton
from google.cloud import speech_v1 as speech1 # need this version for confidence
from google.cloud import storage

######################
# IO functions
######################
def get_sess_audio(sesspath):
    sessname = os.path.basename(sesspath)

    USE_LINKED_MEDIA = True
    # check for linked media first, then for audio files
    if os.path.exists(os.path.join(sesspath, 'LINKED_MEDIA.txt')   ):
        with open(os.path.join(sesspath, 'LINKED_MEDIA.txt')) as lf:
            audiofile = lf.read()
        if os.path.exists(audiofile):
            print(f'...Linked media found: {audiofile}')
        else:
            print(f'...Linked media not found! Will look for media in session directory...')
            USE_LINKED_MEDIA = False
    else:
        USE_LINKED_MEDIA = False

    if not USE_LINKED_MEDIA:
        # prefer wav if it exists, otherwise choose another audio file
        if os.path.exists(os.path.join(sesspath, f'{sessname}.wav')   ):
            audiofile = os.path.join(sesspath, f'{sessname}.wav')   
            print('...local WAV file found.')
        else:
            audiofiles = [f for f in os.listdir(sesspath) if f.split('.')[-1] in ['MOV', 'mov', 'WAV', 'wav', 'mp4', 'mp3', 'm4a', 'aac', 'flac', 'alac', 'ogg']]
            if audiofiles:
                print(3)
                if len(audiofiles) > 1: # choose one format to proceed with
                    for f in audiofiles:
                        if f.split('.')[-1] in ['wav', 'WAV']:
                            audiofile = os.path.join(sesspath, f)
                            continue
                        else:
                            audiofile = os.path.join(sesspath, f)
                else: # only 1 audio file found
                    audiofile = audiofiles[0]
                    aud_type = Path(audiofile).suffix
                print(f'...local {aud_type} file found.')
            else:
                print('!!!WARNING: no audio files found.')
                audiofile=None
    return(audiofile)

######################
# blocking, segmenting, VAD and TAD functions
######################
# Read .wav file, and return (PCM audio data, sample rate)
def read_wave(path):
    with contextlib.closing(wave.open(path, 'rb')) as wf:
        num_channels = wf.getnchannels()
        assert num_channels == 1
        sample_width = wf.getsampwidth()
        assert sample_width == 2
        sample_rate = wf.getframerate()
        assert sample_rate in (8000, 16000, 32000, 48000)
        pcm_data = wf.readframes(wf.getnframes())
        return pcm_data, sample_rate

# write .wav file
def write_wave(path, audio, sample_rate):
    with contextlib.closing(wave.open(path, 'wb')) as wf:
        wf.setnchannels(1)
        wf.setsampwidth(2)
        wf.setframerate(sample_rate)
        wf.writeframes(audio)

# frame of audio data
class Frame(object):
    def __init__(self, bytes, timestamp, duration):
        self.bytes = bytes
        self.timestamp = timestamp
        self.duration = duration

# generates frames of requested duration from audio data
def frame_generator(frame_duration_ms, audio, sample_rate):
    n = int(sample_rate * (frame_duration_ms / 1000.0) * 2)
    offset = 0
    timestamp = 0.0
    duration = (float(n) / sample_rate) / 2.0
    while offset + n < len(audio):
        yield Frame(audio[offset:offset + n], timestamp, duration)
        timestamp += duration
        offset += n

def vad_collector(sample_rate, frame_duration_ms,
                  padding_duration_ms, vad, frames, min_seg_dur = 2000):
    # Filter out non-voiced audio frames.
    # Uses a padded sliding window (ring buffer) over the audio frames.
    # Trigger when more than 90% of the frames in the window are voiced
    # Detrigger when 90% of the frames in the window are unvoiced
    # pads with small amount of silence at start and end
    # Arguments:
    # sample_rate - The audio sample rate, in Hz.
    # frame_duration_ms - The frame duration in milliseconds.
    # padding_duration_ms - The amount to pad the window, in milliseconds.
    # vad - An instance of webrtcvad.Vad.
    # frames - a source of audio frames (sequence or generator).
    # min_seg_dur - minimum segment duration
    # use a deque for sliding window/ring buffer.
    # Parts of this pipeline based on py-webrtcvad The MIT License (MIT) Copyright (c) 2016 John Wiseman
    seg_info = []
    num_padding_frames = int(padding_duration_ms / frame_duration_ms)
    ring_buffer = collections.deque(maxlen=num_padding_frames)

    #  start in the NOTTRIGGERED state.
    triggered = False

    # for segment map - will be stored in tuple seg_info
    seg_start = ''
    seg_end = ''
    numseg = 0

    voiced_frames = []
    for frame in frames:
        is_speech = vad.is_speech(frame.bytes, sample_rate)

        if not triggered:
            ring_buffer.append((frame, is_speech))
            # If more than 90% of the frames in buffer are voiced, TRIGGER
            num_voiced = len([f for f, speech in ring_buffer if speech])
            if num_voiced > 0.9 * ring_buffer.maxlen:
                triggered = True
                #sys.stdout.write('+(%s)' % (ring_buffer[0].timestamp,))
                #sys.stdout.write('+(%s)' % (ring_buffer[0][0].timestamp,))

                # yield all audio from now until NOTTRIGGERED
		        # start with audio already in the ring buffer.
                for f, s in ring_buffer:
                    voiced_frames.append(f)
                seg_start = ring_buffer[0][0].timestamp
                ring_buffer.clear()
        else:
            # in the TRIGGERED state, add data to ring buffer.
            voiced_frames.append(frame)
            ring_buffer.append((frame, is_speech))
            num_unvoiced = len([f for f, speech in ring_buffer if not speech])
            # If more than 90% of the frames in the ring buffer are
            # unvoiced, then enter NOTTRIGGERED and yield whatever
            # audio we've collected.
            if (num_unvoiced > 0.9 * ring_buffer.maxlen) and (len(voiced_frames) >= min_seg_dur/frame_duration_ms):
                triggered = False

                seg_end = frame.timestamp
                seg_info= (numseg, seg_start, seg_end) 
                numseg += 1

                yield (b''.join([f.bytes for f in voiced_frames]) , seg_info)
                ring_buffer.clear()
                voiced_frames = []

    # If we have any leftover voiced audio when we run out of input,
    # yield it.
    if voiced_frames:
        seg_end = frame.timestamp
        seg_info= (numseg, seg_start, seg_end) 
        yield (b''.join([f.bytes for f in voiced_frames]) , seg_info)
        numseg += 1

def segment_coverage_legacy(segfile, sesswav):
    # computes % coverage of VAD-detected segments relative to the original full session file
    # old version using .seg files
    sess_duration = AudioSegment.from_file(sesswav).duration_seconds
    segmap = pd.read_csv(segfile, 
    sep='\s+', header=None, index_col=False, 
    dtype={0:'int',1:'float',2:'float'},
    names = ['segment','start_s','end_s'])   
    segmap['duration'] = segmap['end_s'] - segmap['start_s']
    seg_total_dur = segmap.sum()['duration']
    coverage = seg_total_dur/sess_duration
    return coverage

def segment_coverage(blkfile, sesswav):
    # computes % coverage of VAD-detected segments relative to the original full session file
    # old version using .seg files
    sess_duration = AudioSegment.from_file(sesswav).duration_seconds
    blkmap = pd.read_csv(blkfile, 
    sep='\s+', header=None, index_col=False, 
    dtype={0:'int',1:'int',2:'float',3:'float'},
    names = ['block','segment','start_s','end_s']) 
    blkmap['duration'] = blkmap['end_s'] - blkmap['start_s']
    seg_total_dur = blkmap.sum()['duration']
    coverage = seg_total_dur/sess_duration
    return coverage

######################
# ASR functions
######################

def transcribe_bytestream(bytes, client,srate):
    """Transcribe the given audio bytestream using Google cloud speech."""

    audio = speech.RecognitionAudio(content=bytes)
    config = speech.RecognitionConfig(
        encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
        sample_rate_hertz=srate,
        language_code="en-US",
        model="video"
    )
    result=[]
    try:
        response = client.recognize(config=config, audio=audio)

    # Each result is for a consecutive portion of the audio. Iterate through
    # them to get the transcripts for the entire audio file.
        for r in response.results:
            # The first alternative is the most likely one for this portion.
            best = r.alternatives[0].transcript
            result.append(best)
            print(best)
    except Exception as ex:
        print(f"An exception of type {type(ex).__name__} occurred.")
        raise ex

    return('\n'.join(result))

def transcribeExtra_bytestream(bytes, client,srate):
    """Transcribe the given audio bytestream using Google cloud speech, returning full output 
    including confidence, timing, diarization, alternatives"""

    audio = speech1.RecognitionAudio(content=bytes)
    config = speech1.RecognitionConfig(
        encoding=speech1.RecognitionConfig.AudioEncoding.LINEAR16,
        sample_rate_hertz=srate,
        enable_word_time_offsets=True,
        language_code="en-US",
        use_enhanced=True,
        model="video",
        enable_word_confidence = True,
        max_alternatives = 30)
    result=[]
        
    response = client.recognize(config=config, audio=audio)

    # Each result is for a consecutive portion of the audio. Iterate through
    # them to get the transcripts for the entire audio file.
    best = ''
    for r in response.results:
        # The first alternative is the most likely one for this portion.
        best = r.alternatives[0].transcript
        result.append(best)
        print(best)


    return(response, best)

def transcribe_short_bytestream(bytes, client,srate):
    """Transcribe the given audio bytestream using Google cloud speech-to-text 
    streaming recognize, optimized for short utterances"""
    # see https://cloud.google.com/speech-to-text/docs/best-practices 

    #audio = speech.RecognitionAudio(content=bytes)
    stream = [bytes]
    
    config = speech.RecognitionConfig(
        encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
        sample_rate_hertz=srate,
        language_code="en-US")
    config = speech.StreamingRecognitionConfig(config=config, single_utterance = True
    )   

    requests = (speech.StreamingRecognizeRequest(audio_content=chunk) for chunk in stream )
    result=[]
    try:
        responses = client.streaming_recognize(config=config, requests=requests)

        # Each result is for a consecutive portion of the audio. Iterate through
        # them to get the transcripts for the entire audio file.
        for response in responses:
            for r in response.results:
                # The first alternative is the most likely one for this portion.
                best = r.alternatives[0].transcript
                result.append(best)
                print(best)
    except Exception as ex:
        print(f"An exception of type {type(ex).__name__} occurred.")
        raise ex

    return('\n'.join(result))
    
def transcribe_diarize_file_async(speech_uri, client):
    """Transcribe the given audio file using Google cloud speech,
    with diarization and utterance timings. For best results, run
    on long-duration audio (e.g. >5mins)"""

    audio =speechB.RecognitionAudio(uri=speech_uri)
    config = speechB.RecognitionConfig(
        encoding=speechB.RecognitionConfig.AudioEncoding.LINEAR16,
        sample_rate_hertz=16000,
        language_code="en-US",
        enable_speaker_diarization=True,
        enable_word_time_offsets=True,
        diarization_speaker_count=6,
        use_enhanced=True,
        model="video"    )

    operation = client.long_running_recognize(config=config, audio=audio)
    print('Waiting for recognition to complete...')
    response = operation.result(timeout=3600) # timeout in seconds, default is too short
    result = response.results[-1] # the final element contains the actual transcript
    best = result.alternatives[0].words # choose most likely result
    
    # loop over each word and format the transcript
    transcript=[]
    speaker_tags = []
    speaker_last = None
    words = []
    for w in best:
        words.append({'start_time' : w.start_time.total_seconds(),
        'end_time' : w.end_time.total_seconds(),
        'speaker_tag' : w.speaker_tag,
        'word' : w.word})
        if (speaker_last == w.speaker_tag):
            transcript.append(w.word)
        else: 
            transcript.append(f"\n{w.start_time.total_seconds()}s (speaker {w.speaker_tag}): {w.word}")
        speaker_tags.append(w.speaker_tag)
        speaker_last = w.speaker_tag
    print(' '.join(transcript))

    return transcript, words

def create_bucket(bucket_name, storage_client):
    """Create a new bucket in specific location with storage class"""
    # bucket_name = "your-new-bucket-name"

    bucket = storage_client.bucket(bucket_name)
    if not bucket.exists():
        bucket.storage_class = "STANDARD"
        bucket = storage_client.create_bucket(bucket, location="us")
        print(f"Created bucket {bucket.name} in {bucket.location} with storage class {bucket.storage_class}")
    else:
        print(f"Bucket {bucket_name} already existed")
    return bucket

def upload_blob(source_file_name, bucket_name, destination_blob_name, storage_client):
    """Uploads a file to the bucket."""
    # The ID of your GCS bucket
    # bucket_name = "your-bucket-name"
    # The path to your file to upload
    # source_file_name = "local/path/to/file"
    # The ID of your GCS object
    # destination_blob_name = "storage-object-name"
    
    bucket = storage_client.bucket(bucket_name)
    blob = bucket.blob(destination_blob_name)
    print(f"Uploading {source_file_name} - please wait.")

    ## For slow upload speed
    storage.blob._DEFAULT_CHUNKSIZE = 2097152 # 10242 MB
    storage.blob._MAX_MULTIPART_SIZE = 2097152 # 2 MB
    blob.upload_from_filename(source_file_name,timeout=600.0)

    print(f"File {source_file_name} uploaded.")

######################
# text formatting and WER functions
######################

def strip_punct(instr):
    newstr = ''
    for word in instr.split():
	    # delete punct
        word = word.strip(string.punctuation)

        # delete commas inside numbers
        m = re.match(r'(\d*),(\d)', word)
        if m != None:
            word = word.replace(',', '')

        # commas inside words become space
        word = re.sub(",", " ", word)

        # hyphens inside words become space
        word = re.sub("-", " ", word)
        word = word.strip()

        newstr += ' ' + word
    newstr = newstr.strip()
    return newstr

def remove_in_brackets(text):
    # removes any clause in brackets or parens, and the brackets themselves
    return re.sub("[\(\[].*?[\)\]]+", " ", text)

def caught_num2words(str):
    try:
        return num2words(str)
    except (InvalidOperation, ValueError) as error:
        return str

def format_text_for_wer(text):
    # function to format text or lists of text (e.g. asr, transcript) for wer computation. 
    # Converts from list to a single string and apply some text normalization operations
    # note that the clean_REV_transcript function should be applied first to remove REV-specific keywords 
    # and extract text from docx format tables

    
    if isinstance(text,list):
        text = ' '.join(text)
    text = text.replace('\n',' ') # replace newline with space
    text = remove_in_brackets(text) # removes non-spoken annotations such as [inaudible]
    text = re.sub('%\w+','', text) # remove %HESITATION etc
    text = ' '.join([caught_num2words(str) for str in text.split(' ')]) # spell out numbers
    text = strip_punct(text)
    text = text.lower()
    text = re.sub('\s+',' ',text) # replace multiple space with single

    return text

def clean_REV_transcript(docx_fname, txt_fname):
    doc = docx.Document(docx_fname)
    doctext = [p.text for p in doc.paragraphs]
    # The transcript may be packed into tables
    for t in doc.tables:
        for c in t._cells:
            doctext.append(c.text)

    # write unstripped transcript to .txt
    txt_fname_diarized = re.sub('.txt','_diarized.txt',txt_fname)
    with open(txt_fname_diarized,'w') as outfile:
        outfile.write('\n'.join(doctext))

    # strip the various Speaker IDs and crosstalk indicators  off
    doc_stripped = [re.sub('Speaker \d+:','',line) for line in doctext]
    doc_stripped = [re.sub('.+:','',line) for line in doc_stripped] # remove anything before colon - this is speaker ID
    doc_stripped = [re.sub(r"\t",'',line) for line in doc_stripped] # remove tabs
    doc_stripped = [line  for line in doc_stripped if not re.match(r'^\s*$', line)] # remove blank lines
    doc_stripped = [remove_in_brackets(line) for line in doc_stripped] # remove sections in brackets or parens
    doc_stripped = [strip_punct(line)  for line in doc_stripped] # remove punct
    # write stripped transcript to txt
    with open(txt_fname,'w') as outfile:
        outfile.write('\n'.join(doc_stripped))

def HHMMSS_to_sec(time_str):
    """Get Seconds from time with milliseconds."""
    if time_str.count(':')==2:
        h, m, s = time_str.split(':')
    else:
        print(f'input string format not supported: {time_str}')
    return int(h) * 3600 + int(m) * 60 + float(s) 

def name_counter(str):
    patterns = ['\[redacted[\w\s]*\]', '\[Student \d+\]+']
    N = 0
    for p in patterns:
        N+=len(re.findall(p, str,re.IGNORECASE))
    return N



def align_words(ref,hyp):
    '''
    Aligns two lists of words and outputs the alignment and edit operations
        Parameters:
            ref: reference string
            hyp: hypothesis string


        Returns:
            aligned: a pandas dataframe representing the alignment, 1 row per word 
                with columns:
                    ref_ix: index of word in the reference 
                    hyp_ix: index of word in the hypothesis
                    reference: word from the reference
                    hypothesis: matched word in hypothesis
                    operation: symbolic representations of operation 
                        {'=' : match, 
                        '+':insertion,
                        '-' : deletion,
                        '<>' : substitution
                        }
                    index_edit_ops: index into the edit_ops variable pertaining to each row 
            edit_ops: data frame of word-level operations to go from ref -> hyp

    
    '''

    # get all words and encode as UTF-8 characters to get alignment operations at word-level
    lexicon = list(set(ref+hyp))
    word2digit = dict((lexicon[i],chr(i)) for i in range(0,len(lexicon)))
    asr_uni =  [word2digit[w] for w in hyp]
    trans_uni =  [word2digit[w] for w in ref]
    edit_ops = pd.DataFrame(Levenshtein.editops(''.join(trans_uni),''.join(asr_uni)),
        columns = ['operation','transcript_ix','asr_ix'])
    

    # align the sequences, starting with a dumb alignment where they start together, then inserting as necessary
    aligned_ref = ref.copy()
    aligned_hyp = hyp.copy()
    ix_edit_ops = [np.NaN] *len(aligned_ref)
    aligned_ops =['='] *len(aligned_ref)
    aligned_ref_ix = list(range(len(ref)))
    aligned_hyp_ix = list(range(len(hyp)))

    ins_count = 0 # counter for insertion operations which increase the length of the ref seq thus change the indices
    del_count = 0 # counter for deletion operations which increase the length of the hyp seq thus change the indices
    for [i,ops] in edit_ops.iterrows():
        if ops['operation'] == 'insert':
            aligned_ref.insert(ins_count+ops['transcript_ix'],'_')
            aligned_ops.insert(ins_count+ops['transcript_ix'],'ins')
            aligned_ref_ix.insert(ins_count+ops['transcript_ix'],None)
            ix_edit_ops.insert(ins_count+ops['transcript_ix'],i)
            ins_count = ins_count+1

        if ops['operation'] == 'delete':
            aligned_hyp.insert(del_count+ops['asr_ix'],'_')
            aligned_ops[ins_count + ops['transcript_ix']] = 'del'
            aligned_hyp_ix.insert(del_count+ops['asr_ix'],None)
            ix_edit_ops[ins_count + ops['transcript_ix']] = i
            del_count=del_count+1

        if ops['operation'] == 'replace':
            aligned_ops[ins_count+ ops['transcript_ix']] ='sub' 
            ix_edit_ops[ins_count+ ops['transcript_ix']] =i
           

    aligned = pd.DataFrame({
        'ref_ix':[int(x) if x else None for x in aligned_ref_ix ],
        'hyp_ix':[int(x) if x else None for x in aligned_hyp_ix ],
        'reference':aligned_ref,
        'hypothesis' : aligned_hyp ,
        'operation' : aligned_ops,
        'index_edit_ops' : ix_edit_ops})

    return aligned, edit_ops

def wer_from_counts(N, sub_count, del_count, ins_count):
    '''
    Computes WER and related measures from edit operation counts and reference wordcount
    Useful to recompute measures without needing raw text
    '''
    meas = {'wer': (sub_count + del_count + ins_count)/N,
            'mer': 1 - (N - sub_count - del_count)/N,
            'hits': N - sub_count - del_count ,
            'sub_rate': sub_count/N,
            'del_rate': del_count/N,
            'ins_rate': ins_count/N
            }
    return meas
